import base64
import difflib
import html
import io
import json
from pathlib import Path
import markdown as markdown_lib
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import Response
from sqlalchemy import select
from sqlalchemy.orm import Session
from weasyprint import HTML
from ...database import get_db
from ...config import WORKSPACE_DIR
from ...models import (Evidence, EvidenceImageEdit, Finding, FindingAsset,
    FindingEvidence, FindingRetest,
    ExploitLocalRun, ExploitModification, ExploitResearch, ExploitSource,
    Project, Report, Service, Target)
from ...schemas import ReportIn, ReportOut
from ...time import utcnow

router = APIRouter(prefix="/api/reports", tags=["Reports"])
OSCP_TEMPLATE = """# OSCP Penetration Test Report

## Executive Summary

Write your own summary and assessment.

## Scope

Record the authorized targets and limitations.

## Methodology

Document the manual methodology used during the examination.

## Findings

### Target

#### Finding title

**Observation**

**Impact (user-authored)**

**Reproduction steps**

**Evidence**

## Additional Information

## Evidence Index
"""


def need(db: Session, model, ident: int):
    row = db.get(model, ident)
    if not row:
        raise HTTPException(404, "Not found")
    return row


def values(body: ReportIn) -> dict:
    result = body.model_dump()
    result["evidence_links"] = json.dumps(
        result["evidence_links"], ensure_ascii=False)
    result["exploit_research_links"] = json.dumps(
        list(dict.fromkeys(result["exploit_research_links"])))
    if result["template"] == "oscp" and not result["markdown"]:
        result["markdown"] = OSCP_TEMPLATE
    return result


@router.get("", response_model=list[ReportOut])
def reports(project_id: int, db: Session = Depends(get_db)):
    return db.scalars(select(Report).where(
        Report.project_id == project_id).order_by(Report.id.desc())).all()


@router.post("", response_model=ReportOut, status_code=201)
def create_report(body: ReportIn, db: Session = Depends(get_db)):
    need(db, Project, body.project_id)
    row = Report(**values(body))
    db.add(row); db.commit(); db.refresh(row)
    return row


@router.put("/{ident}", response_model=ReportOut)
def update_report(ident: int, body: ReportIn,
                  db: Session = Depends(get_db)):
    row = need(db, Report, ident)
    if row.project_id != body.project_id:
        raise HTTPException(400, "Project cannot be changed")
    for key, value in values(body).items():
        setattr(row, key, value)
    row.updated_at = utcnow()
    db.commit(); db.refresh(row)
    return row


def render_report(db: Session, row: Report, profile: str = "internal") -> str:
    if profile not in ("client", "internal"):
        raise HTTPException(400, "Profile must be client or internal")
    links = json.loads(row.evidence_links)
    evidence_rows = []
    sensitive = []
    for link in links:
        evidence = need(db, Evidence, int(link["id"]))
        if evidence.project_id != row.project_id:
            raise HTTPException(400, "Linked evidence belongs to another project")
        if profile == "client" and evidence.sensitivity != "normal":
            continue
        if evidence.sensitivity != "normal":
            sensitive.append(evidence.id)
        caption = html.escape(str(link.get("caption", evidence.title)))
        preview = ""
        path = Path(evidence.file_path)
        if evidence.kind == "screenshot" and path.is_file() and path.stat().st_size <= 10_000_000:
            encoded = base64.b64encode(path.read_bytes()).decode()
            preview = f'<img src="data:image/png;base64,{encoded}" alt="{caption}">'
        evidence_rows.append(
            f"<article><h3>Evidence #{evidence.id}: {html.escape(evidence.title)}</h3>"
            f"{preview}<p>{caption}</p><code>SHA-256: {evidence.sha256}</code></article>")
    if sensitive and not row.sensitivity_reviewed:
        raise HTTPException(
            409, f"Sensitive evidence review required for IDs: {sensitive}")
    safe_markdown = html.escape(row.markdown)
    content = markdown_lib.markdown(
        safe_markdown, extensions=["fenced_code", "tables"])
    research_rows = []
    status_labels = {
        "unverified": "미확인", "researching": "조사 중",
        "possibly_affected": "영향 가능", "prerequisites_unmet": "전제조건 불충족",
        "verified": "검증 완료", "not_applicable": "해당 없음",
    }
    for research_id in ([] if profile == "client" else json.loads(row.exploit_research_links or "[]")):
        research = need(db, ExploitResearch, int(research_id))
        if research.project_id != row.project_id:
            raise HTTPException(400, "Linked research belongs to another project")
        modifications = db.scalars(select(ExploitModification).where(
            ExploitModification.research_id == research.id)).all()
        sources = db.scalars(select(ExploitSource).where(
            ExploitSource.research_id == research.id)).all()
        local_runs = db.scalars(select(ExploitLocalRun).where(
            ExploitLocalRun.research_id == research.id).order_by(
            ExploitLocalRun.id)).all()
        changes = "".join(
            f"<li>{html.escape(item.variable_name)}: "
            f"{'••••••••' if item.sensitive else html.escape(item.modified_value)}"
            f" — {html.escape(item.reason)}</li>" for item in modifications)
        source_links = "".join(
            f"<li>{html.escape(item.title)} "
            f"{html.escape(item.source_url)}</li>" for item in sources)
        run_rows = "".join(
            f"<li>#{item.id} {html.escape(item.file_name)} "
            f"SHA-256 {html.escape(item.file_sha256)} — process: "
            f"{html.escape(item.status)}, exit: {item.exit_code}, "
            f"user assessment: {html.escape(item.user_result)}</li>"
            for item in local_runs)
        diffs = []
        root = WORKSPACE_DIR.resolve()
        for source in sources:
            original = Path(source.original_path)
            working = Path(source.working_path)
            if not source.original_path or not source.working_path:
                continue
            try:
                before, after = original.resolve(), working.resolve()
                if (root not in before.parents or root not in after.parents
                        or original.is_symlink() or working.is_symlink()
                        or before.stat().st_size > 1_000_000
                        or after.stat().st_size > 1_000_000):
                    continue
                difference = "".join(difflib.unified_diff(
                    before.read_text(errors="replace").splitlines(True),
                    after.read_text(errors="replace").splitlines(True),
                    fromfile=before.name, tofile=after.name))
                if difference:
                    diffs.append(
                        f"<pre>{html.escape(difference[:200_000])}</pre>")
            except (OSError, UnicodeError):
                continue
        candidate = "검증된 취약점" if research.validation_status == "verified" else "조사 후보"
        research_rows.append(
            f"<article><h3>{candidate}: {html.escape(research.title)}</h3>"
            f"<p><b>상태:</b> {status_labels.get(research.validation_status, research.validation_status)}</p>"
            f"<p><b>대상:</b> {html.escape(research.target_address)} "
            f"{research.port}/{html.escape(research.protocol)} "
            f"{html.escape(research.service_name)} {html.escape(research.discovered_version)}</p>"
            f"<p><b>후보:</b> {html.escape(research.cve)} "
            f"EDB-{html.escape(research.exploit_db_id)}</p>"
            f"<p><b>영향 버전:</b> {html.escape(research.affected_versions)}</p>"
            f"<p><b>발견 근거:</b> {html.escape(research.discovery_evidence)}</p>"
            f"<p><b>전제조건:</b> {html.escape(research.runtime_conditions)} "
            f"{html.escape(research.network_conditions)}</p>"
            f"<p><b>실행 명령:</b> <code>{html.escape(research.execution_command)}</code></p>"
            f"<p><b>실행 결과:</b> {html.escape(research.execution_result)}</p>"
            f"<p><b>재현 절차:</b> {html.escape(research.reproduction_steps)}</p>"
            f"<ul>{source_links}{changes}</ul>"
            f"<h4>Approved local runs</h4><ul>{run_rows}</ul>"
            f"{''.join(diffs)}</article>")
    findings = db.scalars(select(Finding).where(Finding.project_id == row.project_id)).all()
    order = {"Critical": 5, "High": 4, "Medium": 3, "Low": 2, "Informational": 1}
    findings = [x for x in findings if profile == "internal" or x.disclosure != "INTERNAL"]
    findings.sort(key=lambda x: (-order[x.final_risk], -float(x.cvss_score), -x.sort_priority, x.id))
    counts = {level: sum(x.final_risk == level for x in findings) for level in order}
    target_count = len({x.target_id for x in findings if x.target_id})
    service_count = len({x.service_id for x in findings if x.service_id})
    finding_rows = []
    summary_rows = []
    for finding in findings:
        assets = db.scalars(select(FindingAsset).where(
            FindingAsset.finding_id == finding.id)).all()
        target_ids = list(dict.fromkeys([x.target_id for x in assets]
            + ([finding.target_id] if finding.target_id else [])))
        targets_for_finding = [db.get(Target, ident) for ident in target_ids]
        target_label = ", ".join(x.ip for x in targets_for_finding if x)
        summary_rows.append(f"<tr><td>{html.escape(finding.final_risk)}</td><td>{html.escape(finding.title)}</td>"
            f"<td>{html.escape(finding.status)}</td><td>{html.escape(target_label)}</td></tr>")
        links = db.scalars(select(FindingEvidence).where(
            FindingEvidence.finding_id == finding.id).order_by(
            FindingEvidence.display_order, FindingEvidence.id)).all()
        evidence_html = []
        for link in links:
            if profile == "client" and not link.include_client: continue
            if profile == "internal" and not link.include_internal: continue
            ev = db.get(Evidence, link.evidence_id) if link.evidence_id else None
            if not ev or ev.project_id != row.project_id: continue
            if profile == "client" and ev.sensitivity != "normal": continue
            caption = html.escape(link.caption or ev.title)
            preview = ""
            try:
                edit = db.scalar(select(EvidenceImageEdit).where(
                    EvidenceImageEdit.evidence_id == ev.id,
                    EvidenceImageEdit.rendered_path != "").order_by(
                    EvidenceImageEdit.version.desc()))
                path = Path(edit.rendered_path if edit else ev.file_path).resolve()
                root = WORKSPACE_DIR.resolve()
                safe_file = (root in path.parents and path.is_file()
                    and not path.is_symlink())
                if safe_file and ev.kind == "screenshot" and path.stat().st_size <= 10_000_000:
                    content = path.read_bytes()
                    mime = ("image/png" if content.startswith(b"\x89PNG\r\n\x1a\n")
                            else "image/jpeg" if content.startswith(b"\xff\xd8\xff")
                            else "")
                    if mime:
                        encoded = base64.b64encode(content).decode()
                        preview = (f'<figure><img src="data:{mime};base64,{encoded}" '
                            f'alt="{caption}"><figcaption>{caption}</figcaption></figure>')
                elif (profile == "internal" and safe_file
                      and ev.kind in ("command_output", "http", "nmap")
                      and path.stat().st_size <= 1_000_000):
                    preview = f"<pre>{html.escape(path.read_text(errors='replace'))}</pre>"
                elif profile == "internal" and ev.markdown:
                    preview = f"<pre>{html.escape(ev.markdown)}</pre>"
            except OSError:
                preview = ""
            evidence_html.append(f"<li><b>{caption}</b> "
                f"<small>SHA-256 {html.escape(ev.sha256)}</small>{preview}</li>")
        internal = (f"<h3>Technical impact</h3><p>{html.escape(finding.technical_impact)}</p>"
            f"<h3>Reproduction</h3><pre>{html.escape(finding.reproduction_steps)}</pre>"
            f"<h3>Internal notes</h3><pre>{html.escape(finding.internal_notes)}</pre>"
            if profile == "internal" else "")
        finding_rows.append(f"<article class='finding'><h2>{html.escape(finding.title)}</h2>"
            f"<p><b>{html.escape(finding.final_risk)}</b> · CVSS {html.escape(finding.cvss_score)} · "
            f"{html.escape(finding.status)}</p><h3>Summary</h3><p>{html.escape(finding.summary)}</p>"
            f"<h3>Description</h3><p>{html.escape(finding.description)}</p>"
            f"<h3>Business impact</h3><p>{html.escape(finding.business_impact)}</p>{internal}"
            f"<h3>Recommendation</h3><p>{html.escape(finding.recommendation)}</p>"
            f"<h3>Evidence</h3><ul>{''.join(evidence_html)}</ul></article>")
    sensitive_banner = ("<p class='warning'>INTERNAL — MAY CONTAIN SENSITIVE INFORMATION</p>"
                        if profile == "internal" else "")
    return f"""<!doctype html><html><head><meta charset="utf-8"><title>{html.escape(row.title)}</title>
<style>body{{font:11pt sans-serif;max-width:900px;margin:40px auto;line-height:1.5}}
code,pre{{font-family:monospace;background:#f2f2f2}}pre{{padding:12px;white-space:pre-wrap}}
img{{max-width:100%;max-height:700px}}article{{break-inside:avoid;border-top:1px solid #bbb;margin-top:20px}}
.finding{{break-before:page}}.warning{{background:#fee;color:#900;padding:10px}}@page{{margin:18mm;@bottom-center{{content:"Page " counter(page)}}}}
table{{border-collapse:collapse;width:100%}}td,th{{border:1px solid #aaa;padding:5px}}</style></head>
<body>{sensitive_banner}<h1>{html.escape(row.title)}</h1>{content}
<section><h2>Executive Summary</h2><p>This assessment identified {len(findings)} findings across
{target_count} target(s) and {service_count} service(s).</p>
<p>{' · '.join(f'{key}: {value}' for key, value in counts.items())}</p></section>
<section><h2>Finding Summary</h2><table><thead><tr><th>Risk</th><th>Finding</th><th>Status</th><th>Target</th></tr></thead>
<tbody>{''.join(summary_rows)}</tbody></table></section>
<section><h2>Finding Details</h2>{''.join(finding_rows)}</section>
<section><h2>Exploit Research</h2>{''.join(research_rows)}</section>
<section><h2>Evidence Index</h2>{''.join(evidence_rows)}</section></body></html>"""


def _add_markdown(document: Document, source: str) -> None:
    in_code = False
    for line in source.splitlines():
        if line.startswith("```"):
            in_code = not in_code
            continue
        if in_code:
            run = document.add_paragraph().add_run(line)
            run.font.name = "Consolas"; run.font.size = Pt(9)
        elif line.startswith("#"):
            level = min(len(line) - len(line.lstrip("#")), 4)
            document.add_heading(line[level:].strip(), level=level)
        elif line.startswith(("- ", "* ")):
            document.add_paragraph(line[2:], style="List Bullet")
        elif line.strip():
            document.add_paragraph(line)


def _evidence_image(db: Session, evidence: Evidence) -> bytes | None:
    edit = db.scalar(select(EvidenceImageEdit).where(
        EvidenceImageEdit.evidence_id == evidence.id,
        EvidenceImageEdit.rendered_path != "").order_by(
        EvidenceImageEdit.version.desc()))
    path = Path(edit.rendered_path if edit else evidence.file_path)
    try:
        resolved, root = path.resolve(), WORKSPACE_DIR.resolve()
        if (root not in resolved.parents or path.is_symlink()
                or not resolved.is_file() or resolved.stat().st_size > 10_000_000):
            return None
        content = resolved.read_bytes()
        if not (content.startswith(b"\x89PNG\r\n\x1a\n")
                or content.startswith(b"\xff\xd8\xff")):
            return None
        return content
    except OSError:
        return None


def render_docx(db: Session, row: Report, profile: str) -> bytes:
    # Reuse the existing report renderer as the single security/profile gate.
    render_report(db, row, profile)
    document = Document()
    section = document.sections[0]
    section.top_margin = section.bottom_margin = Inches(0.7)
    section.left_margin = section.right_margin = Inches(0.8)
    styles = document.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(10.5)

    if profile == "internal":
        warning = document.add_paragraph()
        warning.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = warning.add_run("INTERNAL — MAY CONTAIN SENSITIVE INFORMATION")
        run.bold = True
    document.add_heading(row.title, 0)
    document.add_paragraph(
        f"Editable {profile.title()} penetration test report")
    _add_markdown(document, row.markdown)

    findings = db.scalars(select(Finding).where(
        Finding.project_id == row.project_id)).all()
    findings = [item for item in findings
                if profile == "internal" or item.disclosure != "INTERNAL"]
    order = {"Critical": 5, "High": 4, "Medium": 3,
             "Low": 2, "Informational": 1}
    findings.sort(key=lambda item: (
        -order[item.final_risk], -float(item.cvss_score),
        -item.sort_priority, item.id))

    document.add_heading("Finding Summary", level=1)
    table = document.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    for cell, label in zip(table.rows[0].cells,
                           ("Risk", "Finding", "Status", "Target")):
        cell.text = label
    for finding in findings:
        assets = db.scalars(select(FindingAsset).where(
            FindingAsset.finding_id == finding.id)).all()
        target_ids = list(dict.fromkeys(
            [asset.target_id for asset in assets]
            + ([finding.target_id] if finding.target_id else [])))
        target_label = ", ".join(
            target.ip for ident in target_ids
            if (target := db.get(Target, ident)))
        cells = table.add_row().cells
        for cell, value in zip(cells, (
                finding.final_risk, finding.title,
                finding.status, target_label)):
            cell.text = value

    document.add_heading("Finding Details", level=1)
    for finding in findings:
        document.add_heading(finding.title, level=2)
        document.add_paragraph(
            f"{finding.final_risk} · CVSS {finding.cvss_score} · {finding.status}")
        fields = [
            ("Summary", finding.summary),
            ("Description", finding.description),
            ("Business impact", finding.business_impact),
        ]
        if profile == "internal":
            fields += [
                ("Technical impact", finding.technical_impact),
                ("Reproduction", finding.reproduction_steps),
                ("Internal notes", finding.internal_notes),
            ]
        fields.append(("Recommendation", finding.recommendation))
        for heading, value in fields:
            if value:
                document.add_heading(heading, level=3)
                document.add_paragraph(value)
        document.add_heading("Evidence", level=3)
        links = db.scalars(select(FindingEvidence).where(
            FindingEvidence.finding_id == finding.id).order_by(
            FindingEvidence.display_order, FindingEvidence.id)).all()
        for link in links:
            if profile == "client" and not link.include_client:
                continue
            if profile == "internal" and not link.include_internal:
                continue
            evidence = db.get(Evidence, link.evidence_id) if link.evidence_id else None
            if (not evidence or evidence.project_id != row.project_id
                    or (profile == "client" and evidence.sensitivity != "normal")):
                continue
            caption = link.caption or evidence.title
            document.add_paragraph(
                f"{caption}\nSHA-256: {evidence.sha256}", style="Caption")
            image = _evidence_image(db, evidence) if evidence.kind == "screenshot" else None
            if image:
                document.add_picture(io.BytesIO(image), width=Inches(6.2))
            elif profile == "internal" and evidence.markdown:
                run = document.add_paragraph().add_run(evidence.markdown)
                run.font.name = "Consolas"; run.font.size = Pt(9)

    document.add_heading("Evidence Index", level=1)
    for link in json.loads(row.evidence_links):
        evidence = db.get(Evidence, int(link["id"]))
        if (not evidence or evidence.project_id != row.project_id
                or (profile == "client" and evidence.sensitivity != "normal")):
            continue
        document.add_paragraph(
            f"Evidence #{evidence.id}: {link.get('caption') or evidence.title}\n"
            f"SHA-256: {evidence.sha256}", style="List Bullet")
        image = _evidence_image(db, evidence) if evidence.kind == "screenshot" else None
        if image:
            document.add_picture(io.BytesIO(image), width=Inches(6.2))

    output = io.BytesIO()
    document.save(output)
    return output.getvalue()


@router.get("/{ident}/export")
def export_report(ident: int, format: str = "html",
                  db: Session = Depends(get_db), profile: str = "internal"):
    row = need(db, Report, ident)
    safe_name = "".join(c if c.isalnum() or c in "._-" else "_"
                        for c in row.title).strip("._") or f"report-{row.id}"
    if format == "docx":
        content = render_docx(db, row, profile)
        return Response(content, media_type=
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition":
                     f'attachment; filename="{safe_name}-{profile}.docx"'})
    document = render_report(db, row, profile)
    if format == "markdown":
        return Response(row.markdown, media_type="text/markdown",
                        headers={"Content-Disposition": f'attachment; filename="{safe_name}.md"'})
    if format == "html":
        return Response(document, media_type="text/html",
                        headers={"Content-Disposition":
                                 f'attachment; filename="{safe_name}.html"'})
    if format == "pdf":
        pdf = HTML(string=document).write_pdf()
        return Response(pdf, media_type="application/pdf",
                        headers={"Content-Disposition":
                                 f'attachment; filename="{safe_name}.pdf"'})
    raise HTTPException(400, "Format must be markdown, html, pdf or docx")
